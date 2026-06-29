#!/usr/bin/env python3
"""Build docs/Требования_к_РСМ_конкурс.docx from _requirements_extract.txt for jury acceptance."""

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "_requirements_extract.txt"
OUT = ROOT / "docs" / "Требования_к_РСМ_конкурс.docx"


def main() -> None:
    text = SRC.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Требования к модулю ресурсно-сервисной модели", level=0)
    doc.add_paragraph("Официальное ТЗ конкурса OmniSight РСМ (сгенерировано из репозитория).")

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped[0].isdigit() and "." in stripped[:4]:
            doc.add_heading(stripped, level=2)
        elif stripped.endswith(":") and len(stripped) < 80:
            doc.add_heading(stripped, level=3)
        else:
            doc.add_paragraph(stripped)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
